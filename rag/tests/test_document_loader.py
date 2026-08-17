import pytest
import tempfile
from pathlib import Path
from docx import Document
import pymupdf

from rag.ingestion.document_loader import (
    load_document,
    load_pdf,
    load_docx,
    detect_language_from_path,
)


class TestDetectLanguageFromPath:
    """Test suite for detect_language_from_path function"""

    def test_detect_english_from_path(self):
        """Test detection of English from path"""
        path = Path("data/raw/en/document.pdf")
        language = detect_language_from_path(path)
        assert language == "en"

    def test_detect_tamil_from_path(self):
        """Test detection of Tamil from path"""
        path = Path("data/raw/ta/document.pdf")
        language = detect_language_from_path(path)
        assert language == "ta"

    def test_detect_hindi_from_path(self):
        """Test detection of Hindi from path"""
        path = Path("data/raw/hi/document.pdf")
        language = detect_language_from_path(path)
        assert language == "hi"

    def test_detect_telugu_from_path(self):
        """Test detection of Telugu from path"""
        path = Path("data/raw/te/document.pdf")
        language = detect_language_from_path(path)
        assert language == "te"

    def test_detect_malayalam_from_path(self):
        """Test detection of Malayalam from path"""
        path = Path("data/raw/ml/document.pdf")
        language = detect_language_from_path(path)
        assert language == "ml"

    def test_unknown_language_returns_unknown(self):
        """Test that unknown language returns 'unknown'"""
        path = Path("data/raw/xx/document.pdf")
        language = detect_language_from_path(path)
        assert language == "unknown"

    def test_case_insensitive_language_detection(self):
        """Test case-insensitive language detection"""
        path = Path("data/raw/EN/document.pdf")
        language = detect_language_from_path(path)
        assert language == "en"


class TestLoadDocx:
    """Test suite for load_docx function"""

    def test_load_simple_docx(self):
        """Test loading a simple DOCX document"""
        with tempfile.TemporaryDirectory() as tmpdir:
            # Create a simple DOCX file
            doc_path = Path(tmpdir) / "en" / "test.docx"
            doc_path.parent.mkdir(parents=True, exist_ok=True)

            doc = Document()
            doc.add_paragraph("First paragraph of test document.")
            doc.add_paragraph("Second paragraph with more content.")
            doc.save(str(doc_path))

            # Load the document
            result = load_docx(doc_path)

            assert result["document_id"] == "test"
            assert result["source"] == "test.docx"
            assert result["file_type"] == "docx"
            assert result["language"] == "en"
            assert result["total_pages"] is None
            assert len(result["paragraphs"]) == 2

    def test_load_docx_with_empty_paragraphs(self):
        """Test that empty paragraphs are skipped"""
        with tempfile.TemporaryDirectory() as tmpdir:
            doc_path = Path(tmpdir) / "en" / "test_empty.docx"
            doc_path.parent.mkdir(parents=True, exist_ok=True)

            doc = Document()
            doc.add_paragraph("First paragraph")
            doc.add_paragraph("")  # Empty paragraph
            doc.add_paragraph("Third paragraph")
            doc.save(str(doc_path))

            result = load_docx(doc_path)

            # Empty paragraphs should be filtered out
            assert len(result["paragraphs"]) == 2

    def test_load_docx_detects_language(self):
        """Test that language is detected from path"""
        with tempfile.TemporaryDirectory() as tmpdir:
            doc_path = Path(tmpdir) / "ta" / "test.docx"
            doc_path.parent.mkdir(parents=True, exist_ok=True)

            doc = Document()
            doc.add_paragraph("தமிழ் பாராக்கிரப்")
            doc.save(str(doc_path))

            result = load_docx(doc_path)

            assert result["language"] == "ta"

    def test_load_docx_preserves_text(self):
        """Test that document text is preserved"""
        with tempfile.TemporaryDirectory() as tmpdir:
            doc_path = Path(tmpdir) / "en" / "test.docx"
            doc_path.parent.mkdir(parents=True, exist_ok=True)

            original_text = "This is important content that must be preserved."
            doc = Document()
            doc.add_paragraph(original_text)
            doc.save(str(doc_path))

            result = load_docx(doc_path)

            assert original_text in result["paragraphs"][0]["text"]


class TestLoadPdf:
    """Test suite for load_pdf function"""

    def test_load_simple_pdf(self):
        """Test loading a simple PDF document"""
        with tempfile.TemporaryDirectory() as tmpdir:
            # Create a simple PDF file
            pdf_path = Path(tmpdir) / "en" / "test.pdf"
            pdf_path.parent.mkdir(parents=True, exist_ok=True)

            doc = pymupdf.open()
            page = doc.new_page()
            page.insert_text((50, 50), "First page content with text.")
            doc.save(str(pdf_path))
            doc.close()

            # Load the document
            result = load_pdf(pdf_path)

            assert result["document_id"] == "test"
            assert result["source"] == "test.pdf"
            assert result["file_type"] == "pdf"
            assert result["language"] == "en"
            assert result["total_pages"] >= 1
            assert len(result["pages"]) >= 1

    def test_load_pdf_multiple_pages(self):
        """Test loading PDF with multiple pages"""
        with tempfile.TemporaryDirectory() as tmpdir:
            pdf_path = Path(tmpdir) / "en" / "test_multi.pdf"
            pdf_path.parent.mkdir(parents=True, exist_ok=True)

            doc = pymupdf.open()
            for i in range(3):
                page = doc.new_page()
                page.insert_text((50, 50), f"Page {i+1} content.")
            doc.save(str(pdf_path))
            doc.close()

            result = load_pdf(pdf_path)

            assert result["total_pages"] == 3
            assert len(result["pages"]) == 3

    def test_load_pdf_page_numbers(self):
        """Test that page numbers are correctly tracked"""
        with tempfile.TemporaryDirectory() as tmpdir:
            pdf_path = Path(tmpdir) / "en" / "test_page_nums.pdf"
            pdf_path.parent.mkdir(parents=True, exist_ok=True)

            doc = pymupdf.open()
            for i in range(2):
                page = doc.new_page()
                page.insert_text((50, 50), f"Page {i+1}")
            doc.save(str(pdf_path))
            doc.close()

            result = load_pdf(pdf_path)

            for i, page in enumerate(result["pages"]):
                assert page["page_number"] == i + 1

    def test_load_pdf_detects_language(self):
        """Test that language is detected from path"""
        with tempfile.TemporaryDirectory() as tmpdir:
            pdf_path = Path(tmpdir) / "hi" / "test.pdf"
            pdf_path.parent.mkdir(parents=True, exist_ok=True)

            doc = pymupdf.open()
            page = doc.new_page()
            page.insert_text((50, 50), "हिंदी पाठ")
            doc.save(str(pdf_path))
            doc.close()

            result = load_pdf(pdf_path)

            assert result["language"] == "hi"


class TestLoadDocument:
    """Test suite for load_document function"""

    def test_load_document_file_not_found(self):
        """Test that FileNotFoundError is raised for missing file"""
        with pytest.raises(FileNotFoundError):
            load_document("/nonexistent/path/file.pdf")

    def test_load_document_unsupported_extension(self):
        """Test that unsupported file types raise ValueError"""
        with tempfile.TemporaryDirectory() as tmpdir:
            file_path = Path(tmpdir) / "test.txt"
            file_path.write_text("Some text content")

            with pytest.raises(ValueError, match="Unsupported file type"):
                load_document(str(file_path))

    def test_load_document_with_pdf_extension(self):
        """Test loading a PDF via load_document"""
        with tempfile.TemporaryDirectory() as tmpdir:
            pdf_path = Path(tmpdir) / "en" / "test.pdf"
            pdf_path.parent.mkdir(parents=True, exist_ok=True)

            doc = pymupdf.open()
            page = doc.new_page()
            page.insert_text((50, 50), "Test content")
            doc.save(str(pdf_path))
            doc.close()

            result = load_document(str(pdf_path))

            assert result["file_type"] == "pdf"
            assert result["document_id"] == "test"

    def test_load_document_with_docx_extension(self):
        """Test loading a DOCX via load_document"""
        with tempfile.TemporaryDirectory() as tmpdir:
            docx_path = Path(tmpdir) / "en" / "test.docx"
            docx_path.parent.mkdir(parents=True, exist_ok=True)

            doc = Document()
            doc.add_paragraph("Test paragraph")
            doc.save(str(docx_path))

            result = load_document(str(docx_path))

            assert result["file_type"] == "docx"
            assert result["document_id"] == "test"

    def test_load_document_case_insensitive_extension(self):
        """Test that file extension matching is case-insensitive"""
        with tempfile.TemporaryDirectory() as tmpdir:
            docx_path = Path(tmpdir) / "en" / "test.DOCX"
            docx_path.parent.mkdir(parents=True, exist_ok=True)

            doc = Document()
            doc.add_paragraph("Test")
            doc.save(str(docx_path))

            result = load_document(str(docx_path))

            assert result["file_type"] == "docx"
